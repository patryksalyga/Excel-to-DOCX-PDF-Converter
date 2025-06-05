import tkinter as tk
from tkinter import filedialog, messagebox, colorchooser
import os
import openpyxl
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas as pdf_canvas
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import mm

def convert_row_to_docx(df, output_path, selected_columns, title_font_size, text_font_size, title_color_rgb):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Pt(595.27)
    section.page_height = Pt(841.89)

    # Ustawienie stopki z dynamicznym numerem strony
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    run = paragraph.add_run("Strona ")
    run.font.size = Pt(10)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    for idx, row in df.iterrows():
        doc.add_page_break() if idx > 0 else None
        for col in selected_columns:
            p_title = doc.add_paragraph()
            run_title = p_title.add_run(col)
            run_title.font.bold = True
            run_title.font.size = Pt(title_font_size)
            run_title.font.color.rgb = RGBColor(*title_color_rgb)

            p_value = doc.add_paragraph()
            run_value = p_value.add_run(str(row[col]))
            run_value.font.size = Pt(text_font_size)

    doc.save(output_path)

def convert_row_to_pdf(df, output_path, selected_columns, title_font_size, text_font_size, title_color_rgb):
    pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
    c = pdf_canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    margin = 20 * mm
    line_spacing = text_font_size + 2
    title_spacing = title_font_size + 4
    footer_height = 15 * mm

    page_number = 1

    def draw_footer():
        c.setFont("Arial", 10)
        c.drawRightString(width - margin, footer_height - 5, f"Strona {page_number}")

    y = height - margin

    for idx, row in df.iterrows():
        for col in selected_columns:
            c.setFont("Arial", title_font_size)
            c.setFillColorRGB(*(v / 255 for v in title_color_rgb))

            if y < margin + footer_height + title_spacing:
                draw_footer()
                c.showPage()
                c.setFont("Arial", title_font_size)
                c.setFillColorRGB(*(v / 255 for v in title_color_rgb))
                y = height - margin
                page_number += 1

            c.drawString(margin, y, str(col))
            y -= title_spacing

            c.setFont("Arial", text_font_size)
            c.setFillColorRGB(0, 0, 0)
            text = str(row[col])
            lines = text.split('\n')

            for line in lines:
                if y < margin + footer_height + line_spacing:
                    draw_footer()
                    c.showPage()
                    c.setFont("Arial", text_font_size)
                    y = height - margin
                    page_number += 1
                c.drawString(margin + 20, y, line)
                y -= line_spacing
            y -= 8
        draw_footer()
        page_number += 1
        c.showPage()
        y = height - margin

    draw_footer()
    c.save()

def run_gui():
    def browse_file():
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        excel_path.set(path)

    def browse_output():
        ext = ".pdf" if export_format.get() == "PDF" else ".docx"
        path = filedialog.asksaveasfilename(defaultextension=ext)
        output_path.set(path)

    def update_column_list():
        try:
            wb = openpyxl.load_workbook(excel_path.get())
            ws = wb.active
            headers = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
            for widget in frame_columns.winfo_children():
                widget.destroy()
            for col in headers:
                var = tk.BooleanVar(value=True)
                checkbox = tk.Checkbutton(frame_columns, text=col, variable=var)
                checkbox.var = var
                checkbox.grid(sticky='w')
                column_vars[col] = var
        except Exception as e:
            messagebox.showerror("Błąd", str(e))
    
    def pick_color():
        color_code = colorchooser.askcolor(title="Wybierz kolor tytułów")
        if color_code:
            entry_color.delete(0, tk.END)
            entry_color.insert(0, color_code[1])

    def convert():
        if not os.path.exists(excel_path.get()):
            messagebox.showerror("Błąd", "Nieprawidłowa ścieżka do pliku Excel.")
            return

        try:
            wb = openpyxl.load_workbook(excel_path.get())
            ws = wb.active
            data = [row for row in ws.iter_rows(values_only=True)]
            df = pd.DataFrame(data[1:], columns=data[0])
            df.dropna(how='all', inplace=True)
            df.dropna(axis=1, how='all', inplace=True)
            df.reset_index(drop=True, inplace=True)

            selected_cols = [col for col, var in column_vars.items() if var.get()]
            if not selected_cols:
                messagebox.showerror("Błąd", "Wybierz przynajmniej jedną kolumnę.")
                return

            tfsize = int(entry_title_font.get())
            vfsize = int(entry_text_font.get())
            color = tuple(int(entry_color.get()[i:i+2], 16) for i in (1, 3, 5))

            if export_format.get() == "DOCX":
                convert_row_to_docx(df[selected_cols], output_path.get(), selected_cols, tfsize, vfsize, color)
                messagebox.showinfo("Sukces", f"Plik zapisany jako: {output_path.get()}")
            else:
                convert_row_to_pdf(df[selected_cols], output_path.get(), selected_cols, tfsize, vfsize, color)
                messagebox.showinfo("Sukces", f"Plik zapisany jako: {output_path.get()}")

        except Exception as e:
            messagebox.showerror("Błąd", str(e))

    root = tk.Tk()
    root.title("Konwerter Excel do DOCX / PDF")

    excel_path = tk.StringVar()
    output_path = tk.StringVar()
    column_vars = {}
    export_format = tk.StringVar(value="DOCX")

    tk.Label(root, text="Format eksportu:").grid(row=0, column=0, sticky='e')
    tk.OptionMenu(root, export_format, "DOCX", "PDF").grid(row=0, column=1, sticky='w')

    tk.Label(root, text="Plik Excel:").grid(row=1, column=0, sticky='e')
    tk.Entry(root, textvariable=excel_path, width=40).grid(row=1, column=1)
    tk.Button(root, text="Wybierz...", command=browse_file).grid(row=1, column=2)
    tk.Button(root, text="Załaduj kolumny", command=update_column_list).grid(row=1, column=3)

    frame_columns = tk.LabelFrame(root, text="Wybierz kolumny do konwersji")
    frame_columns.grid(row=2, column=0, columnspan=4, pady=10, sticky='we')

    tk.Label(root, text="Zapisz jako:").grid(row=3, column=0, sticky='e')
    tk.Entry(root, textvariable=output_path, width=40).grid(row=3, column=1)
    tk.Button(root, text="Wybierz...", command=browse_output).grid(row=3, column=2)

    tk.Label(root, text="Rozmiar czcionki tytułu:").grid(row=4, column=0, sticky='e')
    entry_title_font = tk.Entry(root)
    entry_title_font.insert(0, "14")
    entry_title_font.grid(row=4, column=1)

    tk.Label(root, text="Rozmiar czcionki tekstu:").grid(row=5, column=0, sticky='e')
    entry_text_font = tk.Entry(root)
    entry_text_font.insert(0, "12")
    entry_text_font.grid(row=5, column=1)

    tk.Label(root, text="Kolor tytułów (#RRGGBB):").grid(row=6, column=0, sticky='e')
    entry_color = tk.Entry(root)
    entry_color.insert(0, "#000000")
    entry_color.grid(row=6, column=1)
    tk.Button(root, text="Wybierz kolor", command=pick_color).grid(row=6, column=2)

    tk.Button(root, text="Konwertuj", command=convert, bg="green", fg="white").grid(row=7, column=1, pady=10)

    root.mainloop()


if __name__ == "__main__":
    run_gui()